"""
import docx
from loaders.file_loader import FileLoader

class DOCXLoader(FileLoader):
    def validate_file(self):
        if not self.file_path.endswith('.docx'):
            raise ValueError("Invalid file type. Expected a DOCX file.")

    def load_file(self):
        return docx.Document(self.file_path)
        """

import docx
from loaders.file_loader import FileLoader

class DOCXLoader(FileLoader):
    def validate_file(self):
        if not self.file_path.endswith('.docx'):
            raise ValueError("Invalid file type. Expected a DOCX file.")

    def load_file(self):
        return docx.Document(self.file_path)

    def extract_urls(self):
        """Extract URLs and their associated text from the DOCX file."""
        urls = []
        doc = self.load_file()

        # Iterate through all paragraphs and runs to extract hyperlink text
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Check if the run contains a hyperlink
                for rel in doc.part.rels.values():
                    if "hyperlink" in rel.reltype:
                        if run.text in rel.target_ref:
                            urls.append({"text": run.text, "url": rel.target_ref})

        return urls

